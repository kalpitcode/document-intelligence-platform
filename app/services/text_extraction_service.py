"""
Text Extraction Service Module
===============================

Multi-format document text extraction supporting PDF (PyMuPDF with OCR fallback),
DOCX (python-docx), TXT, and standalone Image assets (PNG, JPEG).
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Sequence

import docx
import fitz  # PyMuPDF
from PIL import Image
import structlog

from app.services.ocr_service import OCRService

logger = structlog.get_logger(__name__)


@dataclass
class PageTextResult:
    """Page-level text extraction result."""
    page_number: int
    text: str
    is_ocr: bool = False
    confidence: float = 100.0


@dataclass
class ExtractionResult:
    """Document-level text extraction result payload."""
    raw_text: str
    page_count: int
    pages: list[PageTextResult]


class TextExtractionService:
    """Service handling multi-format document text extraction."""

    def __init__(self, ocr_service: OCRService | None = None) -> None:
        self.ocr_service = ocr_service or OCRService()

    def extract_text(
        self,
        content: bytes,
        mime_type: str,
        filename: str = "",
    ) -> ExtractionResult:
        """
        Extract text content based on detected MIME type / file extension.

        Args:
            content: Raw binary document content.
            mime_type: MIME type of document asset.
            filename: Document original filename.

        Returns:
            ExtractionResult object containing aggregated raw_text, page_count, and per-page details.
        """
        lower_mime = mime_type.lower()
        lower_filename = filename.lower()

        if lower_mime == "application/pdf" or lower_filename.endswith(".pdf"):
            return self._extract_from_pdf(content)
        elif (
            lower_mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword")
            or lower_filename.endswith(".docx")
        ):
            return self._extract_from_docx(content)
        elif lower_mime.startswith("image/") or any(lower_filename.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
            return self._extract_from_image(content)
        elif lower_mime.startswith("text/") or lower_filename.endswith((".txt", ".csv", ".json", ".xml", ".md", ".html")):
            return self._extract_from_txt(content)
        else:
            # Fallback to UTF-8 decoding
            return self._extract_from_txt(content)

    def _extract_from_pdf(self, content: bytes) -> ExtractionResult:
        """
        Extract text from PDF using PyMuPDF (fitz).
        Falls back to OCR per page if selectable text count is below threshold (< 20 chars).
        """
        page_results: list[PageTextResult] = []
        full_text_parts: list[str] = []

        doc = fitz.open(stream=content, filetype="pdf")
        page_count = len(doc)

        for page_idx in range(page_count):
            page_num = page_idx + 1
            page = doc[page_idx]
            extracted_text = page.get_text("text").strip()

            if len(extracted_text) >= 20:
                # Direct selectable text available
                page_results.append(
                    PageTextResult(
                        page_number=page_num,
                        text=extracted_text,
                        is_ocr=False,
                        confidence=100.0,
                    )
                )
                full_text_parts.append(extracted_text)
            else:
                # Scanned or image-based page - render pixmap and run OCR
                logger.info("Executing OCR fallback for PDF page", page_number=page_num)
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")

                ocr_res = self.ocr_service.extract_text_from_image(img_bytes, page_number=page_num)
                page_results.append(
                    PageTextResult(
                        page_number=page_num,
                        text=ocr_res.text,
                        is_ocr=True,
                        confidence=ocr_res.confidence,
                    )
                )
                full_text_parts.append(ocr_res.text)

        doc.close()
        aggregated_text = "\n\n".join(full_text_parts)
        return ExtractionResult(
            raw_text=aggregated_text,
            page_count=max(1, page_count),
            pages=page_results,
        )

    def _extract_from_docx(self, content: bytes) -> ExtractionResult:
        """Extract text paragraphs and headings from DOCX using python-docx."""
        doc_stream = io.BytesIO(content)
        document = docx.Document(doc_stream)

        paragraphs: list[str] = []
        for para in document.paragraphs:
            clean_p = para.text.strip()
            if clean_p:
                paragraphs.append(clean_p)

        # Include tables content in docx text stream if any
        for table in document.tables:
            for row in table.rows:
                row_str = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_str:
                    paragraphs.append(row_str)

        full_text = "\n\n".join(paragraphs)
        page_res = PageTextResult(page_number=1, text=full_text, is_ocr=False, confidence=100.0)

        return ExtractionResult(
            raw_text=full_text,
            page_count=1,
            pages=[page_res],
        )

    def _extract_from_image(self, content: bytes) -> ExtractionResult:
        """Extract text from image asset via OCR Service."""
        ocr_res = self.ocr_service.extract_text_from_image(content, page_number=1)
        page_res = PageTextResult(
            page_number=1,
            text=ocr_res.text,
            is_ocr=True,
            confidence=ocr_res.confidence,
        )
        return ExtractionResult(
            raw_text=ocr_res.text,
            page_count=1,
            pages=[page_res],
        )

    def _extract_from_txt(self, content: bytes) -> ExtractionResult:
        """Read plain text content with UTF-8 / latin-1 encoding fallbacks."""
        text = ""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        page_res = PageTextResult(page_number=1, text=text, is_ocr=False, confidence=100.0)
        return ExtractionResult(
            raw_text=text,
            page_count=1,
            pages=[page_res],
        )
